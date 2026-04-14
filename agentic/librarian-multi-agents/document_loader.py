"""
Document Loader Utilities
Handles loading and extracting text from various document formats
"""

import os
from typing import Optional, Dict, Any
from pathlib import Path

class DocumentLoader:
    """
    Utility class for loading documents from various formats
    Supports: PDF, DOCX, TXT, MD, HTML
    """
    
    @staticmethod
    def load_text_file(file_path: str) -> str:
        """
        Load plain text file
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content as string
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    @staticmethod
    def load_pdf(file_path: str) -> Dict[str, Any]:
        """
        Load PDF file and extract text
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with content, page_count, and metadata
        """
        try:
            import PyPDF2 # you need to install PyPDF2 if not already installed
        except ImportError:
            raise ImportError(
                "PyPDF2 is required for PDF processing. "
                "Install it with: pip install PyPDF2 --break-system-packages"
            )
        
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            
            pages_content = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                pages_content.append({
                    "page_number": page_num,
                    "content": text
                })
            
            # Combine all pages with page markers
            full_content = "\n\n".join([
                f"--- Page {page['page_number']} ---\n{page['content']}"
                for page in pages_content
            ])
            
            return {
                "content": full_content,
                "page_count": len(pdf_reader.pages),
                "pages": pages_content,
                "metadata": pdf_reader.metadata or {}
            }
    
    @staticmethod
    def load_docx(file_path: str) -> Dict[str, Any]:
        """
        Load DOCX file and extract text
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with content, paragraph_count, and metadata
        """
        try:
            from docx import Document # you need to install python-docx if not already installed
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install it with: pip install python-docx --break-system-packages"
            )
        
        doc = Document(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        full_content = "\n\n".join(paragraphs)
        
        # Extract tables if present
        tables_content = []
        for table_num, table in enumerate(doc.tables, 1):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables_content.append({
                "table_number": table_num,
                "data": table_data
            })
        
        return {
            "content": full_content,
            "paragraph_count": len(paragraphs),
            "tables": tables_content,
            "metadata": {
                "core_properties": {
                    "author": doc.core_properties.author,
                    "created": doc.core_properties.created,
                    "modified": doc.core_properties.modified,
                    "title": doc.core_properties.title
                }
            }
        }
    
    @staticmethod
    def load_markdown(file_path: str) -> str:
        """
        Load Markdown file
        
        Args:
            file_path: Path to markdown file
            
        Returns:
            File content as string
        """
        return DocumentLoader.load_text_file(file_path)
    
    @staticmethod
    def load_html(file_path: str) -> str:
        """
        Load HTML file and optionally extract text
        
        Args:
            file_path: Path to HTML file
            
        Returns:
            File content as string
        """
        return DocumentLoader.load_text_file(file_path)
    
    @staticmethod
    def load_document(file_path: str, extract_metadata: bool = True) -> Dict[str, Any]:
        """
        Auto-detect file type and load document
        
        Args:
            file_path: Path to document
            extract_metadata: Whether to extract metadata
            
        Returns:
            Dictionary with content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        loaders = {
            '.txt': DocumentLoader.load_text_file,
            '.md': DocumentLoader.load_markdown,
            '.pdf': DocumentLoader.load_pdf,
            '.docx': DocumentLoader.load_docx,
            '.html': DocumentLoader.load_html,
            '.htm': DocumentLoader.load_html
        }
        
        if extension not in loaders:
            raise ValueError(
                f"Unsupported file format: {extension}. "
                f"Supported formats: {', '.join(loaders.keys())}"
            )
        
        loader = loaders[extension]
        
        # Load the document
        if extension in ['.pdf', '.docx']:
            result = loader(str(file_path))
            return result
        else:
            content = loader(str(file_path))
            return {
                "content": content,
                "file_type": extension,
                "file_size": file_path.stat().st_size
            }

class DocumentSaver:
    """
    Utility class for saving processed documents
    """
    
    @staticmethod
    def save_text(content: str, file_path: str):
        """Save content as text file"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    @staticmethod
    def save_markdown(content: str, file_path: str):
        """Save content as markdown file"""
        DocumentSaver.save_text(content, file_path)
    
    @staticmethod
    def save_html(content: str, file_path: str, title: str = "Document"):
        """
        Save content as HTML file with basic styling
        
        Args:
            content: Content to save
            file_path: Output file path
            title: HTML document title
        """
        html_template = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1, h2, h3 {{ color: #2c3e50; }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #3498db;
            color: white;
        }}
        tr:nth-child(even) {{
            background-color: #f2f2f2;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
    </style>
</head>
<body>
    {content}
</body>
</html>"""
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_template)
    
    @staticmethod
    def save_to_docx(content: str, file_path: str, title: str = "Document"):
        """
        Save content as DOCX file
        
        Args:
            content: Content to save
            file_path: Output file path
            title: Document title
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX creation. "
                "Install it with: pip install python-docx --break-system-packages"
            )
        
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content (split by double newlines for paragraphs)
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
        
        doc.save(file_path)

# Usage example
if __name__ == "__main__":
    print("Document Loader Utilities")
    print("=" * 50)
    
    # Example 1: Load a text file
    print("\nExample 1: Loading text file")
    sample_text = """
    SAMPLE DOCUMENT
    
    This is a sample document for testing the document loader.
    It contains multiple paragraphs and sections.
    
    SECTION 1
    First section content here.
    
    SECTION 2
    Second section content here.
    """
    
    # Save sample
    with open('/tmp/sample.txt', 'w') as f:
        f.write(sample_text)
    
    # Load it back
    loader = DocumentLoader()
    result = loader.load_document('/tmp/sample.txt')
    print(f"Loaded {len(result['content'])} characters")
    
    # Example 2: Save as different formats
    print("\nExample 2: Saving in different formats")
    
    saver = DocumentSaver()
    
    # Save as markdown
    saver.save_markdown(sample_text, '/tmp/output.md')
    print("Saved as markdown: /tmp/output.md")
    
    # Save as HTML
    saver.save_html(sample_text, '/tmp/output.html', title="Sample Document")
    print("Saved as HTML: /tmp/output.html")
    
    print("\nNote: For PDF and DOCX support, install additional packages:")
    print("  pip install PyPDF2 python-docx --break-system-packages")
